#!/usr/bin/env python3
# ─────────────────────────────────────────────────────────────────────────────
# Strip leftover guidance content from a pre-filled PRD. Run this AFTER
# generate-prd.js and BEFORE self-audit.js / present_files.
#
# generate-prd.js emits three kinds of italic flag text, because the same
# file also produces the blank template:
#   ✏  guidance     — "how to fill out this section" notes for the author.
#                      These must NEVER reach a pre-filled, stakeholder-facing
#                      draft.
#   ⚑  Gap: ...      — actual flags that content wasn't in the source material.
#                      These MUST stay visible in the delivered document.
#   ⚠  Conflict: ... — flags that the source material had CONTRADICTORY
#                      values for something (Changeset 3C). These MUST also
#                      stay visible in the delivered document, exactly like
#                      gap flags — only guidance (✏) is ever stripped.
#
# Guidance text shows up in TWO structurally different places, and both must
# be handled — this is the incident that motivated the second half of this
# script: several tables (KPIs, Requirements, User Stories, Edge Cases,
# Constraints, Stakeholders, Decisions, Version History) have a hardcoded
# trailing row of guidance cells baked into the table itself. python-docx's
# `document.paragraphs` does NOT descend into table cells, so a paragraph-only
# strip silently leaves these rows in place — the document looks clean on a
# naive check but still ships guidance content inside tables.
#
#   1. Standalone body paragraphs containing "✏" (outside any table) — removed
#      as whole paragraphs.
#   2. Table rows where ANY cell contains "✏" — removed as whole rows, since
#      every such row in this skill's tables is a dedicated all-guidance
#      trailer row (never a row that mixes real data with guidance in a way
#      that would lose content by removing the row).
#
# Usage:  python3 scripts/strip-guidance.py /path/to/prd-output.docx
# Overwrites the file in place and prints a before/after count.
# ─────────────────────────────────────────────────────────────────────────────
import sys
import docx

if len(sys.argv) != 2:
    print("Usage: python3 scripts/strip-guidance.py /path/to/prd-output.docx", file=sys.stderr)
    sys.exit(2)

path = sys.argv[1]


def iter_all_tables(doc):
    """Yield every table in the document, including tables nested inside cells."""
    def walk(tables):
        for t in tables:
            yield t
            for row in t.rows:
                for cell in row.cells:
                    yield from walk(cell.tables)
    yield from walk(doc.tables)


def full_text_counts(doc):
    """Count '✏', '⚑', and '⚠' occurrences across body paragraphs AND table cells."""
    guidance = sum(1 for p in doc.paragraphs if "✏" in p.text)
    gaps = sum(1 for p in doc.paragraphs if "⚑" in p.text)
    conflicts = sum(1 for p in doc.paragraphs if "⚠" in p.text)
    for t in iter_all_tables(doc):
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "✏" in p.text:
                        guidance += 1
                    if "⚑" in p.text:
                        gaps += 1
                    if "⚠" in p.text:
                        conflicts += 1
    return guidance, gaps, conflicts


d = docx.Document(path)
guidance_before, gaps_before, conflicts_before = full_text_counts(d)

# 1. Remove standalone body paragraphs containing "✏"
body = d.element.body
for p in list(d.paragraphs):
    if "✏" in p.text:
        body.remove(p._element)

# 2. Remove table rows where any cell contains "✏"
for t in iter_all_tables(d):
    for row in list(t.rows):
        if any("✏" in cell.text for cell in row.cells):
            t._tbl.remove(row._tr)

d.save(path)

d2 = docx.Document(path)
guidance_after, gaps_after, conflicts_after = full_text_counts(d2)

print(f"Guidance content removed: {guidance_before} -> {guidance_after}")
print(f"Gap flags preserved: {gaps_before} -> {gaps_after}")
print(f"Conflict flags preserved: {conflicts_before} -> {conflicts_after}")

if guidance_after != 0:
    print("FAILED — guidance content still present after stripping (check nested tables).", file=sys.stderr)
    sys.exit(1)
if gaps_after != gaps_before:
    print("FAILED — gap flag count changed; a gap may have been deleted by mistake.", file=sys.stderr)
    sys.exit(1)
if conflicts_after != conflicts_before:
    print("FAILED — conflict flag count changed; a conflict may have been deleted by mistake.", file=sys.stderr)
    sys.exit(1)

print("Strip-guidance passed.")
