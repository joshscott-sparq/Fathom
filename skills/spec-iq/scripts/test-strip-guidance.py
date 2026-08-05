#!/usr/bin/env python3
# ─────────────────────────────────────────────────────────────────────────────
# Fixture test for scripts/strip-guidance.py (Changeset 3C).
#
# This repo has no other automated tests — all prior testing has been the
# manual "run a full PRD through the pipeline" pass documented in README.md's
# version history. This script is new test infrastructure, added specifically
# because docs/scorecard-proposal.md §E calls out that ⚠ survival "needs an
# explicit test" now that strip-guidance.py has a third marker to get right.
#
# It builds a minimal .docx with python-docx covering three cases, deliberately
# NOT mixed in the same row (production tables never mix them — guidance rows
# are always dedicated all-✏ trailer rows, never a row that also carries real
# ⚑/⚠ content):
#   1. A standalone body paragraph containing "✏"        -> must be REMOVED.
#   2. A standalone body paragraph containing "⚠"        -> must SURVIVE.
#   3. A table row with "⚠" in a data cell, no "✏"
#      anywhere in that row                               -> row must SURVIVE
#                                                             intact.
#   4. A table row where a cell contains "✏" (a dedicated
#      guidance trailer row, matching how generate-prd.js
#      actually emits these)                               -> row must be
#                                                             REMOVED.
# A standalone "⚑" paragraph is included too, exercising the existing gap
# preservation alongside the new conflict preservation.
#
# strip-guidance.py is a top-level script (sys.exit calls, not an importable
# module), so it's invoked here as a subprocess, exactly how the real
# pipeline calls it.
#
# Usage:  python3 scripts/test-strip-guidance.py
# Exits 0 and prints "All assertions passed." on success; exits 1 on any
# failed assertion.
# ─────────────────────────────────────────────────────────────────────────────
import subprocess
import sys
import tempfile
import os

import docx

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
STRIP_GUIDANCE = os.path.join(SCRIPT_DIR, "strip-guidance.py")

failures = []


def check(condition, message):
    if not condition:
        failures.append(message)


with tempfile.TemporaryDirectory() as tmpdir:
    path = os.path.join(tmpdir, "fixture.docx")

    d = docx.Document()
    d.add_paragraph("✏  Guidance: this whole paragraph must be stripped")
    d.add_paragraph("⚑  Gap: this whole paragraph must survive")
    d.add_paragraph("⚠  Conflict: this whole paragraph must survive")

    table = d.add_table(rows=2, cols=2)
    # Row 0: a real data row that happens to carry a conflict flag in one
    # cell — no "✏" anywhere in this row, so it must survive intact.
    table.rows[0].cells[0].text = "Some real content"
    table.rows[0].cells[1].text = "⚠  Conflict: two values disagree"
    # Row 1: a dedicated guidance trailer row (matches how generate-prd.js's
    # guidanceCell() emits these) — must be removed as a whole row.
    table.rows[1].cells[0].text = "✏  Fill this column in with..."
    table.rows[1].cells[1].text = "✏  Guidance text for this row"

    d.save(path)

    # Everything below must stay INSIDE this `with` block — the temp
    # directory (and fixture.docx with it) is deleted the moment the block
    # exits, so strip-guidance.py's subprocess and the re-read afterward both
    # need the file to still be on disk.
    result = subprocess.run(
        [sys.executable, STRIP_GUIDANCE, path],
        capture_output=True,
        text=True,
    )
    print(result.stdout)
    if result.stderr:
        print(result.stderr, file=sys.stderr)

    check(result.returncode == 0, f"strip-guidance.py exited {result.returncode}, expected 0")

    d2 = docx.Document(path)
    body_text = "\n".join(p.text for p in d2.paragraphs)

    check("✏" not in body_text, "guidance marker (✏) survived in a standalone paragraph")
    check("Gap: this whole paragraph must survive" in body_text, "gap flag (⚑) paragraph was removed")
    check("Conflict: this whole paragraph must survive" in body_text, "conflict flag (⚠) standalone paragraph was removed")

    check(len(d2.tables) == 1, "expected exactly one table to survive")
    if d2.tables:
        rows = d2.tables[0].rows
        check(len(rows) == 1, f"expected exactly 1 table row to survive (the data row), found {len(rows)}")
        if rows:
            row_text = " ".join(c.text for c in rows[0].cells)
            check("⚠" in row_text, "conflict flag (⚠) in a table data cell did not survive")
            check("Some real content" in row_text, "the surviving row lost its real data content")
            check("✏" not in row_text, "a guidance trailer row was not removed")

if failures:
    print("FAILED:", file=sys.stderr)
    for f in failures:
        print(f"  - {f}", file=sys.stderr)
    sys.exit(1)

print("All assertions passed.")
